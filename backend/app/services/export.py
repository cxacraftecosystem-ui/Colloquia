"""Transcript / conversation export to TXT, Markdown, DOCX and PDF.

The exported document is the refined CONVERSATION (the cleaned, speaker-labelled dialogue) — the same
thing the app shows on the Conversation tab — rendered with real formatting: bold speaker labels,
headings, bullets and section rules. When a recording was translated, the English version is appended
as its own section. Everything is defensive so a single odd character can never 500 an export.
"""

import io
import re
from typing import Any


def _fmt_ms(ms: int) -> str:
    total = int((ms or 0) / 1000)
    h, rem = divmod(total, 3600)
    m, s = divmod(rem, 60)
    return f"{h:02d}:{m:02d}:{s:02d}" if h else f"{m:02d}:{s:02d}"


def _segments_markdown(segments: list[Any]) -> str:
    lines = []
    for s in segments:
        speaker = f"**{s.speaker}:** " if s.speaker else ""
        lines.append(f"{speaker}{s.editedText or s.text}")
    return "\n\n".join(lines)


def _conversation_sections(recording: Any, transcript: Any, segments: list[Any]) -> list[tuple[str, str]]:
    """Ordered (heading, markdown) sections to render. Prefers the refined dialogue; appends the
    English translation when present; falls back to segments then raw text."""
    sections: list[tuple[str, str]] = []
    refined = getattr(transcript, "refinedText", None) if transcript else None
    translated = getattr(transcript, "translatedText", None) if transcript else None
    if refined and refined.strip():
        sections.append(("Conversation", refined.strip()))
    if translated and translated.strip():
        sections.append(("English translation", translated.strip()))
    if not sections and segments:
        sections.append(("Transcript", _segments_markdown(segments)))
    if not sections and transcript and getattr(transcript, "rawText", None):
        sections.append(("Transcript", transcript.rawText.strip()))
    return sections


def _is_rule(line: str) -> bool:
    return len(line) >= 3 and set(line) <= {"-", "*", "_"}


_INLINE_TOKEN = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*|_(.+?)_|`(.+?)`")


# --------------------------------------------------------------------------- plain text / markdown

def _strip_md(text: str) -> str:
    return re.sub(r"\*\*(.+?)\*\*", r"\1", re.sub(r"`(.+?)`", r"\1", text)).replace("**", "")


def build_txt(recording: Any, transcript: Any, segments: list[Any]) -> bytes:
    title = recording.title or "Recording"
    out = [title, "=" * len(title), ""]
    for heading, md in _conversation_sections(recording, transcript, segments):
        out.append(heading.upper())
        out.append("")
        for line in md.replace("\r\n", "\n").split("\n"):
            out.append(_strip_md(line))
        out.append("")
    return ("\n".join(out) + "\n").encode("utf-8")


def build_markdown(recording: Any, transcript: Any, segments: list[Any]) -> bytes:
    title = recording.title or "Recording"
    out = [f"# {title}", ""]
    sections = _conversation_sections(recording, transcript, segments)
    multi = len(sections) > 1
    for heading, md in sections:
        if multi:
            out.append(f"## {heading}")
            out.append("")
        out.append(md)
        out.append("")
    return ("\n".join(out) + "\n").encode("utf-8")


# --------------------------------------------------------------------------- DOCX

def _docx_inline_runs(paragraph: Any, text: str) -> None:
    """Add `text` to a python-docx paragraph, honouring **bold**, *italic*/_italic_ and `code`."""
    pos = 0
    for m in _INLINE_TOKEN.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        bold, ital_a, ital_b, code = m.groups()
        if bold is not None:
            paragraph.add_run(bold).bold = True
        elif ital_a is not None:
            paragraph.add_run(ital_a).italic = True
        elif ital_b is not None:
            paragraph.add_run(ital_b).italic = True
        elif code is not None:
            run = paragraph.add_run(code)
            run.font.name = "Consolas"
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def build_docx(recording: Any, transcript: Any, segments: list[Any]) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(recording.title or "Recording", level=0)
    sections = _conversation_sections(recording, transcript, segments)
    multi = len(sections) > 1
    for heading, md in sections:
        if multi:
            doc.add_heading(heading, level=1)
        for raw in md.replace("\r\n", "\n").split("\n"):
            line = raw.strip()
            if not line:
                continue
            try:
                if _is_rule(line):
                    doc.add_paragraph("—" * 24)
                elif line.startswith("### "):
                    doc.add_heading(_strip_md(line[4:]), level=3)
                elif line.startswith("## "):
                    doc.add_heading(_strip_md(line[3:]), level=2)
                elif line.startswith("# "):
                    doc.add_heading(_strip_md(line[2:]), level=2)
                elif line.startswith("- ") or line.startswith("* "):
                    _docx_inline_runs(doc.add_paragraph(style="List Bullet"), line[2:])
                elif re.match(r"^\d+\. ", line):
                    _docx_inline_runs(doc.add_paragraph(style="List Number"), line[line.index(".") + 2:])
                else:
                    _docx_inline_runs(doc.add_paragraph(), line)
            except Exception:  # noqa: BLE001 - never let one odd line break the export
                doc.add_paragraph(_strip_md(line))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --------------------------------------------------------------------------- PDF

def _pdf_escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _pdf_inline_html(text: str) -> str:
    """Markdown inline -> reportlab mini-HTML (<b>/<i>/<font>)."""
    out: list[str] = []
    pos = 0
    for m in _INLINE_TOKEN.finditer(text):
        if m.start() > pos:
            out.append(_pdf_escape(text[pos:m.start()]))
        bold, ital_a, ital_b, code = m.groups()
        if bold is not None:
            out.append(f"<b>{_pdf_escape(bold)}</b>")
        elif ital_a is not None:
            out.append(f"<i>{_pdf_escape(ital_a)}</i>")
        elif ital_b is not None:
            out.append(f"<i>{_pdf_escape(ital_b)}</i>")
        elif code is not None:
            out.append(f"<font face='Courier'>{_pdf_escape(code)}</font>")
        pos = m.end()
    if pos < len(text):
        out.append(_pdf_escape(text[pos:]))
    return "".join(out)


def build_pdf(recording: Any, transcript: Any, segments: list[Any]) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, title=recording.title or "Recording")
    styles = getSampleStyleSheet()
    flow: list[Any] = [Paragraph(_pdf_escape(recording.title or "Recording"), styles["Title"]), Spacer(1, 12)]
    sections = _conversation_sections(recording, transcript, segments)
    multi = len(sections) > 1
    for heading, md in sections:
        if multi:
            flow.append(Paragraph(_pdf_escape(heading), styles["Heading1"]))
            flow.append(Spacer(1, 6))
        for raw in md.replace("\r\n", "\n").split("\n"):
            line = raw.strip()
            if not line:
                flow.append(Spacer(1, 4))
                continue
            try:
                if _is_rule(line):
                    flow.append(HRFlowable(width="100%", thickness=0.5, spaceBefore=6, spaceAfter=6))
                elif line.startswith("### "):
                    flow.append(Paragraph(_pdf_inline_html(line[4:]), styles["Heading3"]))
                elif line.startswith("## "):
                    flow.append(Paragraph(_pdf_inline_html(line[3:]), styles["Heading2"]))
                elif line.startswith("# "):
                    flow.append(Paragraph(_pdf_inline_html(line[2:]), styles["Heading2"]))
                elif line.startswith("- ") or line.startswith("* "):
                    flow.append(Paragraph("• " + _pdf_inline_html(line[2:]), styles["BodyText"]))
                else:
                    flow.append(Paragraph(_pdf_inline_html(line), styles["BodyText"]))
                    flow.append(Spacer(1, 4))
            except Exception:  # noqa: BLE001
                flow.append(Paragraph(_pdf_escape(_strip_md(line)), styles["BodyText"]))
    doc.build(flow)
    return buf.getvalue()


EXPORTERS = {
    "txt": ("text/plain", build_txt),
    "md": ("text/markdown", build_markdown),
    "docx": ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", build_docx),
    "pdf": ("application/pdf", build_pdf),
}
